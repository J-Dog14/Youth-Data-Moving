from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import glob

def safe_add_picture(paragraph, image_path, width, height):
    """
    Attempts to add a picture to a Word document paragraph.
    If the image file does not exist, it skips adding the image.
    """
    if os.path.exists(image_path):
        paragraph.add_run().add_picture(image_path, width=width, height=height)
    else:
        print(f"Warning: Image file not found at {image_path}, skipping.")

def create_formatted_word_doc(image_paths, doc_path):
    doc = Document()

    # Footplant
    doc.add_heading('Footplant', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    safe_add_picture(cell.paragraphs[0], image_paths.get('Front@FP', ''), Inches(3.21), Inches(3.21))
    cell = table.cell(0, 1)
    safe_add_picture(cell.paragraphs[0], image_paths.get('sag@FP', ''), Inches(3.21), Inches(3.21))
    doc.add_paragraph()

    # Max External Rotation and Ball Release
    heading = doc.add_heading('', level=2)
    run = heading.add_run('Max External Rotation' + ' ' * 20 + 'Ball Release')
    run.bold = True
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    safe_add_picture(cell.paragraphs[0], image_paths.get('sag@MaxER', ''), Inches(3.21), Inches(3.21))
    cell = table.cell(0, 1)
    safe_add_picture(cell.paragraphs[0], image_paths.get('sag@Rel', ''), Inches(3.21), Inches(3.21))
    doc.add_paragraph()

    # Kinematics Report
    doc.add_heading('Kinematics Report', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if 'Action+Report' in image_paths:
        safe_add_picture(doc.add_paragraph(), image_paths['Action+Report'], Inches(6), Inches(4.5))

    # Athletic Screen Report
    doc.add_heading('Athletic Screen Report', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if 'Athletic_Screen_Report' in image_paths:
        safe_add_picture(doc.add_paragraph(), image_paths['Athletic_Screen_Report'], Inches(6), Inches(4.5))

    # Save the document
    doc.save(doc_path)



# Example usage
image_paths = {
    'Front@FP': 'D:\\Action+\\Exports\\Front@FP.png',
    'sag@FP': 'D:\\Action+\\Exports\\sag@FP.png',
    'sag@MaxER': 'D:\\Action+\\Exports\\sag@MaxER.png',
    'sag@Rel': 'D:\\Action+\\Exports\\sag@Rel.png',
    'Action+Report': 'D:\\Action+\\Exports\\Action+Report.png',
    'Athletic_Screen_Report': 'D:\\Action+\\Exports\\Athletic_Screen_Report.png',
}
doc_path = 'G:\\My Drive\\Action+ Reports\\New Player Data.docx'

create_formatted_word_doc(image_paths, doc_path)

def clear_images_from_folder(folder_path):
    # Define the pattern for PNG files
    pattern = os.path.join(folder_path, '*.png')

    # Use glob to find all files matching the pattern
    files = glob.glob(pattern)

    # Loop through the found files and delete them
    for file in files:
        try:
            os.remove(file)
            print(f"Deleted {file}")
        except OSError as e:
            print(f"Error: {e.strerror} - {file}")


# Example usage
source_folder = 'D:\\Action+\\Exports'
clear_images_from_folder(source_folder)
